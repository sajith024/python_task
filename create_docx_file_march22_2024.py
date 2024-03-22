from docx import Document
from docx.shared import Inches, Cm
import pypandoc


records = [
    {
        "name": "Joe Root",
        "description": "Joseph Edward Root, MBE is an English international cricketer, who plays for the English cricket team and formerly captained the Test team. He also represents Yorkshire in English domestic cricket.",
        "run_score": 10629,
    },
    {
        "name": "Steve Smith",
        "description": "Steven Peter Devereux Smith is an Australian international cricketer and former captain of the Australian national team in all three formats of the game. He is widely regarded as one of the best Test batsmen of all time since Don Bradman.",
        "run_score": 8647,
    },
    {
        "name": "David Warner",
        "description": "David Andrew Warner is an Australian T20I international cricketer and a former Test vice-captain. A left-handed opening batsman, Warner is the first Australian cricketer in 132 years to be selected for the national team in any format without experience in first-class cricket.",
        "run_score": 8132,
    },
    {
        "name": "Virat Kohli",
        "description": "Virat Kohli is an Indian international cricketer and the former captain of the Indian national cricket team. He is a right-handed batsman and an occasional medium-fast bowler. He currently represents Royal Challengers Bengaluru in the IPL and Delhi in domestic cricket.",
        "run_score": 8119,
    },
    {
        "name": "Kane Williamson",
        "description": "Kane Stuart Williamson is a New Zealand international cricketer and captain of the New Zealand national team in limited overs cricket. On 27 February 2023, Williamson became the all-time leading run-scorer for New Zealand in Test cricket.",
        "run_score": 7645,
    },
    {
        "name": "Azhar Ali",
        "description": "Azhar Ali is a Pakistani former international cricketer. He is former captain of ODI and test side of Pakistan national team. Ali made his Test debut for Pakistan against Australia in the first Test at Lord's in July 2010.Azhar Ali is a Pakistani former international cricketer. He is former captain of ODI and test side of Pakistan national team. Ali made his Test debut for Pakistan against Australia in the first Test at Lord's in July 2010.",
        "run_score": 7142,
    },
    {
        "name": "Cheteshwar Pujara",
        "description": "Cheteshwar Arvind Pujara is an Indian cricketer and is the captain of Sussex County Cricket Club in County Championship. He plays for Saurashtra in Indian domestic cricket. Pujara is known for his disciplined and patient batting style which made him an integral part of the Indian Test team for over a decade.",
        "run_score": 7014,
    },
    {
        "name": "Angelo Mathews",
        "description": "Angelo Davis Mathews is a professional Sri Lankan cricketer and a former captain of the national cricket team in all formats.",
        "run_score": 6953,
    },
    {
        "name": "Dimuth Karunaratne",
        "description": "Frank Dimuth Madushanka Karunaratne, popularly known as Dimuth Karunaratne, is a professional Sri Lankan cricketer and former captain of the Sri Lanka Test and ODI teams. Considered as one of the best Test openers in the world, he has included 3 times in the ICC Test Cricket Team of the Year. ",
        "run_score": 6023,
    },
    {
        "name": "Ben Stokes",
        "description": "Benjamin Andrew Stokes OBE is an English international cricketer who is the captain of the England Test team and plays for the England team in ODIs and T20Is. In domestic cricket, he represents Durham and has played in multiple Twenty20 leagues around the world.",
        "run_score": 5602,
    },
]

document = Document()

document.add_heading("Task: Docx creation", 0)

document.add_heading("Batsman Run Score Table", level=1)
table = document.add_table(rows=1, cols=4)
table.style = "Table Grid"

hdr_cells = table.rows[0].cells

paragraph = hdr_cells[0].add_paragraph()
run = paragraph.add_run("Name")
run.bold = True

paragraph = hdr_cells[1].add_paragraph()
run = paragraph.add_run("Description")
run.bold = True

paragraph = hdr_cells[2].add_paragraph()
run = paragraph.add_run("Image")
run.bold = True

paragraph = hdr_cells[3].add_paragraph()
run = paragraph.add_run("Runs Scored")
run.bold = True

table.columns[0].width = Cm(3.5)
table.columns[1].width = Cm(6)
table.columns[2].width = Cm(4.5)
table.columns[3].width = Cm(2)

for bats_man in records:
    row_cells = table.add_row().cells
    row_cells[1].text = bats_man["description"]
    row_cells[3].text = str(bats_man["run_score"])

    paragraph = row_cells[0].paragraphs[0]
    run = paragraph.add_run(bats_man["name"])
    run.bold = True

    paragraph = row_cells[2].paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(
        f'assests/images/{bats_man["name"]}.jpeg',
        width=Inches(1.25),
        height=Inches(1),
    )

document.add_page_break()
document.save("batsman_run.docx")

pypandoc.convert_file("batsman_run.docx", "pdf", outputfile="batsman_run.pdf")
